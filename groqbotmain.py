import os
from groq import Groq
import streamlit as st
import pandas as pd
import openpyxl
from io import StringIO
from PyPDF2 import PdfReader
from docx import Document

from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

os.environ["LANGSMITH_TRACING"]="true"
os.environ["LANGSMITH_TRACING_V2"]="true"
os.environ["LANGSMITH_API_KEY"]=st.secrets["LangSmith_API"]
os.environ["LANGSMITH_PROJECT"]="test"
from langchain_groq import ChatGroq

# Add custom CSS to hide the GitHub icon
# Universal CSS to hide Streamlit clutter
import streamlit as st

hide_everything = """
    <style>
        /* Hide GitHub / Deploy button (top-right corner) */
        [data-testid="stToolbar"] {visibility: hidden !important;}

        /* Hide hamburger menu */
        #MainMenu {visibility: hidden;}

        /* Hide footer */
        footer {visibility: hidden;}

        /* Hide viewer badge (bottom-right, "Made with Streamlit") */
        .viewerBadge_container__1QSob {display: none !important;}

        /* Hide Manage app (bottom-right) */
        [data-testid="stActionButton"] {display: none !important;}
        /* Main background */
        .stApp {
            background-color: #394743;
        }
        /* Sidebar background */
        section[data-testid="stSidebar"] {
            background-color: #72a0c1;
        }
    </style>
"""
st.markdown(hide_everything, unsafe_allow_html=True)

# Initialize chat history with a health-focused system prompt
if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "system",
            "content": (
                "You are a professional Project management office AI assistant. you are specialized in Agile, TOGAF, Prince2, PMI, PMP, Scrum, Spotify, Lean, Kanban, SixSigma and other project management frameworks"
                "Your role is to help users understand their project status and suggest possible common risk, causes and risk rosolution, root cause analysys"
                "you can provide output in excel, word, charts and other presentable ways if asked by user. show output in tabular format. Give detailed analysis and improvement ideas about the project after project status"
                "If a user asks about anything unrelated to Project management, PMO, Project/program, frameworks, Management, risk, resources, reply: "
                "'I'm here to help with PMO related questions and suggestions. "
                "Please ask about project, program, risks, resources or other PMO related questions/concerns.' "
            )
        }
    ]

# Display all previous messages
#for msg in st.session_state.messages[1:]:  # Skip system prompt in UI
#    with st.chat_message(msg["role"]):
#        st.markdown(msg["content"])

# Create prompt template from the stored messages
prompt = ChatPromptTemplate.from_messages([
    (msg["role"], msg["content"]) for msg in st.session_state.messages] + [
    ("user", "{user_input}")
])

#prompt=ChatPromptTemplate.from_messages(
#    [
#        ("system","You are a helpful assistant. Please response to the user queries"),
#        ("user","Question:{question}")
#    ]
#)

llm = ChatGroq(
    groq_api_key=st.secrets["GROQ_API_KEY"],
    model="llama-3.3-70b-versatile",   # or "llama3-70b-8192", etc.
    temperature=0.1  # default is 0.7, change to e.g., 0.5 for more focused output
    #top_p=0.9          
)

#client = Groq(api_key=st.secrets["GROQ_API_KEY"],)
# Set the app title
st.set_page_config(page_title="AI PMO Agent", page_icon="🤖")
st.title("🤖 PMO AI Agent")

# Add some colored headers
#st.markdown("### 🟢 Hi, How can I help you today !!")
#st.text_area("💬 Your Question:", height=100)

# ---- Sidebar = Chat input area ----
with st.sidebar:
    st.header("💬 Chat with Bot")
    with st.form("chat_form", clear_on_submit=True):
        #user_input = st.text_input("Hi, How can I help you today ?", key="input_box")
        user_input=st.text_area("Hi, How can I help you today ?", height=50)
        uploaded_file = st.file_uploader(
        "Attach a supporting file (optional)", 
        type=["txt", "pdf", "docx", "csv","xlsx"]
        )
        submitted = st.form_submit_button("Submit")
    
    #user_input=st.text_input("Hi, How can I help you today ?")
    #Submit button to control execution
    if submitted:
        if not user_input:  # Text is mandatory
            st.error("⚠️ Please enter text before submitting.")
        else:
            st.success("✅ Processing your request...")
            #st.write("Text entered:", text_val)    
            if uploaded_file:
                st.write("File uploaded:", uploaded_file.name)
            else:
                st.info("No file uploaded (that’s okay!)")
                
#text_data=""    
if uploaded_file is not None:
    file_type = uploaded_file.name.split(".")[-1].lower()
    content = ""
    contentTbl=""
    #st.write(file_type)
    if file_type == "txt":
        # Read as text
        stringio = StringIO(uploaded_file.getvalue().decode("utf-8"))
        content = stringio.read()
        text_data=content
        #st.write(text_data)
    elif file_type == "pdf":
        # Extract text from PDF
        reader = PdfReader(uploaded_file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                content += page_text + "\n"
        text_data=content
        #st.write(text_data)
    elif file_type == "docx":
        doc = Document(uploaded_file)
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text is None:
                para.text = ""
                content += para.text + "\n"
            else:
                content += para.text + "\n"
                #content += (para.text or "") + "\n"
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text is None:
                    row_text=""
                    contentTbl += row_text + "\n"
                else:
                    contentTbl += row_text + "\n"
        text_data=content + contentTbl
        #st.write("the text file data is :")
        #st.write(content)
        #st.write("the text file data is :")
        #st.write(contentTbl)
        #st.write(text_data)

    elif file_type == "csv":
        # Read CSV into DataFrame
        df = pd.read_csv(uploaded_file)
        st.dataframe(df)
        text_data = df.to_string()

    elif file_type == "xlsx":
        # Read Excel into DataFrame
        df = pd.read_excel(uploaded_file)
        # Convert to string (you can filter/clean before sending)
        text_data = df.to_string()
    #if content:
    #    st.subheader("Extracted Content:")
    #    st.text(content[:2000])  # Show first 2000 chars (avoid overload)
    else:
        text_data=""
    #Print text file uploaded data
    #st.write(text_data)
    
    # Process user input
    if user_input:
        # Add user's message to history and show
        st.session_state.messages.append({"role": "user", "content": user_input+text_data})
        #with st.chat_message("user")
            #st.markdown(user_input)
        
output_parser=StrOutputParser()
chain=prompt|llm|output_parser
if user_input:
    response=chain.invoke({"user_input":user_input})
    st.markdown(response)

